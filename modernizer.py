import flet as ft
import openai
from openai import OpenAI
import docx
from docx.oxml import OxmlElement
import threading
import time
import os
import json
import sys

if getattr(sys, 'frozen', False):  # If running as a PyInstaller bundle
    base_path = sys._MEIPASS
else:
    base_path = os.path.dirname(__file__)

SETTINGS_FILE = os.path.join(base_path, "settings.json")

# Define constants for theme colors
THEME = {
    "background_color": "#03111f",  # Dark navy blue
    "border_color": "#5697cc",      # Light blue for borders
    "text_color": "#FFFFFF",        # White for text
    "button_color": "#76e3b3",      # Soft green for buttons
    "button_text_color": "#03111f", # Dark navy blue for button text
    "label_color": "#5697cc",       # Light blue for labels
}

def load_settings():
    """Load settings from a file or return default settings."""
    if os.path.exists(SETTINGS_FILE):
        try:
            with open(SETTINGS_FILE, "r") as f:
                return json.load(f)
        except json.JSONDecodeError:
            print("Error: Failed to decode settings.json. Loading default settings.")
    return {
        "api_key": "",
        "model": "gpt-4o",
        "prompt": (
            "You are a helpful assistant that modernizes sentences from old English books into modern English. "
            "Don't use commas too much and change to active voice when possible and if it is just a name, leave it as just the name. "
            "If a Bible verse is quoted, use the ESV Version and give the full name of the book with the reference (example: Deuteronomy 2:12). "
            "I will give you a sentence and you will modernize it without any comment."
        )
    }

settings = load_settings()

client = OpenAI(api_key=settings["api_key"])

def get_available_models():
    try:
        if not settings["api_key"]:
            return ["gpt-4o", "gpt-4", "gpt-3.5-turbo"]
        response = client.models.list()
        models = response.data  # Use `.data` instead of subscriptable access
        return [model.id for model in models]  # Access model IDs correctly
    except Exception as e:
        print(f"Error fetching models: {e}")
        return ["gpt-4o", "gpt-4", "gpt-3.5-turbo"]

def modernize_text(text, max_retries=3, retry_delay=5):
    for attempt in range(max_retries):
        try:
            response = client.chat.completions.create(
                model=settings["model"],
                messages=[
                    {"role": "system", "content": settings["prompt"]},
                    {"role": "user", "content": text}
                ]
            )
            # Access the content attribute directly
            return response.choices[0].message.content
        except Exception as e:
            print(f"Error: {e}")
            if attempt < max_retries - 1:
                time.sleep(retry_delay)
            else:
                raise e




def set_cell_borders(cell, border_width_pt=1):
    cell_properties = cell._element.get_or_add_tcPr()
    borders = ['top', 'bottom', 'left', 'right']
    for border in borders:
        tag = f'w:{border}'
        border_element = OxmlElement(tag)
        border_element.set(docx.oxml.ns.qn('w:val'), 'single')
        border_element.set(docx.oxml.ns.qn('w:sz'), str(border_width_pt * 8))
        border_element.set(docx.oxml.ns.qn('w:space'), '0')
        border_element.set(docx.oxml.ns.qn('w:color'), 'auto')
        cell_properties.append(border_element)

def process_document(input_path, output_path, progress_callback):
    try:
        doc = docx.Document(input_path)
        new_doc = docx.Document()

        # Count non-empty paragraphs
        non_empty_paragraphs = [p for p in doc.paragraphs if p.text.strip() != '']
        total_paragraphs = len(non_empty_paragraphs)
        print(f"Total non-empty paragraphs to process: {total_paragraphs}")

        processed_count = 0  # Track processed paragraphs

        for i, paragraph in enumerate(doc.paragraphs):
            original_text = paragraph.text
            if original_text.strip() == '':
                print(f"Skipping empty paragraph at index {i}.")
                continue

            print(f"Processing paragraph {processed_count + 1}/{total_paragraphs}: {original_text[:50]}...")
            modernized_text = modernize_text(original_text)
            print(f"Modernized paragraph {processed_count + 1}/{total_paragraphs}: {modernized_text[:50]}...")

            # Create a table with two columns
            table = new_doc.add_table(rows=1, cols=2)
            for cell in table.columns[0].cells + table.columns[1].cells:
                set_cell_borders(cell)

            row = table.rows[0]
            row.cells[0].text = original_text
            row.cells[1].text = modernized_text
            new_doc.add_paragraph()

            # Update progress
            processed_count += 1
            progress = processed_count / total_paragraphs
            print(f"Updating progress: {progress * 100:.2f}%")
            progress_callback(progress)

        new_doc.save(output_path)
        print(f"Document saved to {output_path}")
    except Exception as e:
        print(f"Error processing document: {e}")

# Main function
def main(page: ft.Page):
    page.title = "Text Modernizer"
    page.scroll = ft.ScrollMode.AUTO

    # Define colors based on the screenshot
    background_color = THEME["background_color"]
    border_color = THEME["border_color"]  # Light blue for borders and text field outlines
    text_color = THEME["text_color"]  # White for text
    button_color = THEME["button_color"]  # Soft green for buttons
    button_text_color = THEME["button_text_color"]  # Dark navy blue for button text
    label_color = THEME["label_color"]  # Light blue for labels

    # Apply the background color to the page
    page.bgcolor = background_color

    input_file = ft.TextField(
        label="Selected File Path",
        expand=True,
        disabled=True,
        color=text_color,
        bgcolor=background_color,
        border_color=border_color,
        prefix_icon=ft.Icons.FOLDER,  # Add folder icon
    )
    output_file = ft.TextField(
        label="Output File Name (with .docx)",
        expand=True,
        disabled=True,
        color=text_color,
        bgcolor=background_color,
        border_color=border_color,
        prefix_icon=ft.Icons.DESCRIPTION,  # Add document icon
    )
    progress_bar = ft.ProgressBar(width=400, visible=False, bgcolor=background_color, color=border_color)
    progress_label = ft.Text("Progress: 0%", weight="bold", visible=False, color=label_color)

    # FilePicker to replace tkinter's filedialog
    file_picker = ft.FilePicker()
    page.overlay.append(file_picker)

    def update_progress(value):
        progress_bar.value = value
        progress_label.value = f"Progress: {int(value * 100)}%"
        page.update()

    def select_file(e):
        file_picker.pick_files(allowed_extensions=["docx"])

    def file_picked(e: ft.FilePickerResultEvent):
        if e.files:
            input_path = e.files[0].path
            input_file.value = input_path

            # Generate output file name
            output_path = input_path.rsplit(".", 1)[0] + "-modernized.docx"
            output_file.value = output_path

            page.update()

    def open_settings(e):
        def save_and_close_settings(e):
            # Save the settings and close the dialog
            settings["api_key"] = api_key_field.value
            settings["model"] = model_dropdown.value
            settings["prompt"] = prompt_field.value
            with open(SETTINGS_FILE, "w") as f:
                json.dump(settings, f, indent=4)
            settings_dialog.open = False
            page.update()
        
        def cancel_settings(e):
            # Close the dialog without saving
            settings_dialog.open = False
            page.update()
    
        # Define colors based on the theme
        background_color = THEME["background_color"]
        border_color = THEME["border_color"]
        text_color = THEME["text_color"]
        button_color = THEME["button_color"]
        button_text_color = THEME["button_text_color"]
        label_color = THEME["label_color"]
    
        # Settings dialog components
        api_key_field = ft.TextField(
            label="OpenAI API Key",
            value=settings["api_key"],
            password=True,
            expand=True,
            color=text_color,
            bgcolor=background_color,
            border_color=border_color,
        )
        model_dropdown = ft.Dropdown(
            label="ChatGPT Model",
            options=[ft.dropdown.Option(model) for model in get_available_models()],
            value=settings["model"],
            expand=True,
            color=text_color,
            bgcolor=background_color,
            border_color=border_color,
        )
        prompt_field = ft.TextField(
            label="Prompt",
            value=settings["prompt"],
            multiline=True,
            expand=True,
            color=text_color,
            bgcolor=background_color,
            border_color=border_color,
        )
    
        # Save and Cancel buttons
        save_button = ft.ElevatedButton(
            content=ft.Row(
                [
                    ft.Icon(ft.Icons.SAVE, color=THEME["button_text_color"]),
                    ft.Text("Save", color=THEME["button_text_color"]),
                ]
            ),
            on_click=save_and_close_settings,
            style=ft.ButtonStyle(
                bgcolor=THEME["button_color"],
                shape=ft.RoundedRectangleBorder(radius=12),
            ),
        )
        cancel_button = ft.OutlinedButton(
            content=ft.Row(
                [
                    ft.Icon(ft.Icons.CANCEL, color=THEME["border_color"]),
                    ft.Text("Cancel", color=THEME["border_color"]),
                ]
            ),
            on_click=cancel_settings,
            style=ft.ButtonStyle(
                shape=ft.RoundedRectangleBorder(radius=12),
                side=ft.BorderSide(color=THEME["border_color"], width=2),
            ),
        )
    
        # Create the settings dialog
        settings_dialog = ft.AlertDialog(
            title=ft.Text("Settings", color=THEME["text_color"]),
            content=ft.Column(
                [
                    api_key_field,
                    model_dropdown,
                    prompt_field,
                ],
                spacing=15,
                expand=True,
            ),
            actions=[
                ft.Row(
                    [cancel_button, save_button],
                    alignment=ft.MainAxisAlignment.END,
                    spacing=20,
                )
            ],
            modal=True,
            bgcolor=THEME["background_color"],
        )
    
        # Append the dialog to page.overlay
        page.overlay.append(settings_dialog)
        settings_dialog.open = True
        page.update()

    def start_processing(e):
        if not input_file.value or not output_file.value or not settings["api_key"]:
            page.snack_bar = ft.SnackBar(ft.Text("Please fill all fields."))
            page.snack_bar.open = True
            page.update()
            return

        progress_bar.visible = True
        progress_label.visible = True
        page.update()

        threading.Thread(
            target=process_document,
            args=(input_file.value, output_file.value, update_progress),
        ).start()

    file_picker.on_result = file_picked

    # Centered layout with updated styling and icons
    page.add(
        ft.Container(
            content=ft.Column(
                [
                    ft.Row(
                        [
                            ft.ElevatedButton(
                                content=ft.Row(
                                    [
                                        ft.Icon(ft.Icons.UPLOAD_FILE, color=button_text_color),  # Add upload icon
                                        ft.Text("Select File to Modernize", color=button_text_color),
                                    ]
                                ),
                                on_click=select_file,
                                bgcolor=button_color,
                                style=ft.ButtonStyle(shape=ft.RoundedRectangleBorder(radius=12)),
                            ),
                            input_file,
                        ],
                        alignment=ft.MainAxisAlignment.CENTER,
                    ),
                    ft.Row([output_file], alignment=ft.MainAxisAlignment.CENTER),
                    ft.Row(
                        [
                            ft.ElevatedButton(
                                content=ft.Row(
                                    [
                                        ft.Icon(ft.Icons.SETTINGS, color=button_text_color),  # Add settings icon
                                        ft.Text("Settings", color=button_text_color),
                                    ]
                                ),
                                on_click=open_settings,
                                bgcolor=button_color,
                                style=ft.ButtonStyle(shape=ft.RoundedRectangleBorder(radius=12)),
                            ),
                            ft.ElevatedButton(
                                content=ft.Row(
                                    [
                                        ft.Icon(ft.Icons.PLAY_ARROW, color=button_text_color),  # Add play icon
                                        ft.Text("Start Modernization", color=button_text_color),
                                    ]
                                ),
                                on_click=start_processing,
                                bgcolor=button_color,
                                style=ft.ButtonStyle(shape=ft.RoundedRectangleBorder(radius=12)),
                            ),
                        ],
                        alignment=ft.MainAxisAlignment.CENTER,
                        spacing=20,
                    ),
                    ft.Row([progress_bar], alignment=ft.MainAxisAlignment.CENTER),
                    ft.Row([progress_label], alignment=ft.MainAxisAlignment.CENTER),
                ],
                alignment=ft.MainAxisAlignment.CENTER,
                horizontal_alignment=ft.CrossAxisAlignment.CENTER,
                spacing=20,
            ),
            expand=True,  # Make the container fill the window
            alignment=ft.alignment.center,  # Center the content within the container
        )
    )

ft.app(target=main)




